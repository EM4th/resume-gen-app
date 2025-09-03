"""
Clean Flask Resume Generator App with Ad Monetization
"""

from flask import Flask, request, render_template, jsonify, send_file
from flask_cors import CORS
import os
import logging
import uuid
from datetime import datetime
import shutil
import google.generativeai as genai
import pdfplumber
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib import colors
import re
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = Flask(__name__)
CORS(app)

# Configuration
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
UPLOAD_FOLDER = 'uploads'
PREVIEWS_FOLDER = 'previews'

# Create necessary directories
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PREVIEWS_FOLDER, exist_ok=True)

# Configure Gemini AI
GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')
if GOOGLE_API_KEY:
    genai.configure(api_key=GOOGLE_API_KEY)
    model = genai.GenerativeModel('gemini-1.5-flash')
else:
    model = None
    logger.warning("GOOGLE_API_KEY not found. AI features will be disabled.")

# Simple analytics tracking
analytics_data = {
    'daily_users': 0,
    'total_resumes_generated': 0
}

@app.route('/')
def index():
    """Serve the main page with ad integration"""
    analytics_data['daily_users'] += 1
    session_id = str(uuid.uuid4())
    return render_template('index.html', session_id=session_id)

@app.route('/health')
def health_check():
    """Health check endpoint for monitoring"""
    return jsonify({
        'status': 'healthy',
        'timestamp': datetime.now().isoformat(),
        'version': '3.0',
        'app': 'resume-gen'
    })

@app.route('/ads.txt')
def ads_txt():
    """Serve ads.txt file for AdSense verification"""
    try:
        # Try to serve from root directory first
        file_path = 'ads.txt'
        if os.path.exists(file_path):
            return send_file(file_path, mimetype='text/plain')
        # Fallback content
        return "google.com, pub-7524647518323966, DIRECT, f08c47fec0942fa0", 200, {'Content-Type': 'text/plain'}
    except Exception as e:
        logger.error(f"Error serving ads.txt: {str(e)}")
        return "google.com, pub-7524647518323966, DIRECT, f08c47fec0942fa0", 200, {'Content-Type': 'text/plain'}

@app.route('/test')
def test():
    return render_template('test_simple.html')

@app.route('/generate_resume', methods=['POST'])
def generate_resume():
    """Generate resume with AI enhancement"""
    try:
        # Get form data
        job_description_raw = request.form.get('job_description', '').strip()
        output_format = request.form.get('format', 'pdf')

        # Validate input
        if not job_description_raw:
            return jsonify({'error': 'Job description is required'}), 400
        
        # Check if resume file was uploaded
        if 'resume_file' not in request.files:
            return jsonify({'error': 'No resume file provided'}), 400
        
        file = request.files['resume_file']
        if file.filename == '':
            return jsonify({'error': 'No file selected'}), 400
        
        if not file.filename.lower().endswith('.pdf'):
            return jsonify({'error': 'Only PDF files are supported'}), 400
        
        # Save uploaded file
        filename = f"{uuid.uuid4()}_resume.pdf"
        filepath = os.path.join(UPLOAD_FOLDER, filename)
        file.save(filepath)
        
        # Process the resume
        # Expand job description if it's a URL
        if is_probable_url(job_description_raw):
            scraped = fetch_job_content(job_description_raw)
            job_description = scraped if scraped else job_description_raw
            job_source_url = job_description_raw
        else:
            job_description = job_description_raw
            job_source_url = None

        enhanced_resume_path, explanation, extra_files = process_resume_with_ai(filepath, job_description, output_format)
        
        # Update analytics
        analytics_data['total_resumes_generated'] += 1
        
        # Return success with explanation
        resp = {
            'success': True,
            'preview_url': f'/preview/{os.path.basename(enhanced_resume_path)}',
            'download_url': f'/download/{os.path.basename(enhanced_resume_path)}',
            'explanation': explanation,
            'job_url': job_source_url,
            'additional_downloads': {k: f'/download/{os.path.basename(v)}' for k, v in extra_files.items()} if extra_files else {}
        }
        return jsonify(resp)
        
    except Exception as e:
        logger.error(f"Error generating resume: {str(e)}")
        return jsonify({'error': f'Error processing resume: {str(e)}'}), 500

def process_resume_with_ai(resume_path: str, job_description: str, output_format: str) -> tuple[str, str, dict]:
    """Process resume with AI enhancement and return main file path, explanation, and any extra files."""
    try:
        # Extract text from uploaded resume
        resume_text = extract_text_from_pdf(resume_path)
        
        if not resume_text.strip():
            raise ValueError("Could not extract text from PDF")
        
        # Generate enhanced resume with AI
        if model:
            enhanced_content, explanation = enhance_resume_with_ai(resume_text, job_description)
        else:
            # Fallback: basic processing without AI
            enhanced_content = resume_text
            explanation = "AI enhancement unavailable. Original resume returned."
            logger.warning("AI processing disabled, returning original content")
        
        # Generate output file
        output_filename = f"{uuid.uuid4()}_enhanced_resume.pdf"
        output_pdf_path = os.path.join(PREVIEWS_FOLDER, output_filename)
        create_pdf_resume(enhanced_content, output_pdf_path)

        extra = {}
        if output_format.lower() == 'docx':
            docx_filename = f"{uuid.uuid4()}_enhanced_resume.docx"
            docx_path = os.path.join(PREVIEWS_FOLDER, docx_filename)
            create_docx_resume(enhanced_content, docx_path)
            extra['docx'] = docx_path
        elif output_format.lower() == 'pdf':
            # Optionally also create docx for user convenience
            try:
                docx_filename = f"{uuid.uuid4()}_enhanced_resume.docx"
                docx_path = os.path.join(PREVIEWS_FOLDER, docx_filename)
                create_docx_resume(enhanced_content, docx_path)
                extra['docx'] = docx_path
            except Exception as ce:
                logger.warning(f"DOCX generation failed (non-fatal): {ce}")

        return output_pdf_path, explanation, extra
        
    except Exception as e:
        logger.error(f"Error in AI processing: {str(e)}")
        # Fallback: just copy the original file
        output_filename = f"{uuid.uuid4()}_resume.{output_format}"
        output_path = os.path.join(PREVIEWS_FOLDER, output_filename)
        shutil.copy2(resume_path, output_path)
    return output_path, "Error occurred during AI enhancement. Original resume returned.", {}

def extract_text_from_pdf(pdf_path: str) -> str:
    """Extract text from PDF file"""
    try:
        with pdfplumber.open(pdf_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        return ""

#############################################
# AI Enhancer with Robust Formatting & Explain
#############################################

def _normalize_ai_text(text: str) -> str:
    """Normalize AI output to consistent plain text resume format."""
    if not text:
        return ""
    lines = []
    for raw in text.splitlines():
        line = raw.strip('\r ').rstrip()
        if not line:
            lines.append("")
            continue
        # Normalize bullets
        if re.match(r"^[•*\-]\s+", line):
            body = re.sub(r"^[•*\-]\s+", "", line).strip()
            lines.append(f"- {body}")
            continue
        # Uppercase section headers heuristically
        if (len(line) < 40 and 1 <= len(line.split()) <= 5 and
            any(tok.lower() in {"summary","experience","education","skills","projects","certifications","profile","objective"} for tok in line.lower().split())):
            lines.append("")
            lines.append(line.upper())
            lines.append("")
            continue
        lines.append(line)
    # Collapse multiple blank lines
    cleaned = []
    blank = False
    for l in lines:
        if l == "":
            if not blank:
                cleaned.append("")
            blank = True
        else:
            cleaned.append(l)
            blank = False
    return "\n".join(cleaned).strip()


def _heuristic_explanation(original: str, enhanced: str, job: str) -> str:
    """Fallback explanation if model fails or returns unusable text."""
    def tokens(s):
        return set(re.findall(r"[A-Za-z]{3,}", s.lower()))
    o, e, j = tokens(original), tokens(enhanced), tokens(job)
    new_job_terms = (e - o) & j
    improved_overlap = len((e & j)) - len((o & j))
    bullets = [
        f"• **Keyword Alignment:** Added or emphasized terms: {', '.join(list(new_job_terms)[:8]) or 'relevant role-specific keywords' }.",
        f"• **Relevance Increase:** Net gain of ~{improved_overlap if improved_overlap>0 else 0} job-aligned keywords improving ATS match.",
        "• **Action Impact:** Weak verbs replaced with stronger action verbs for clearer ownership and results.",
        "• **Structure & Clarity:** Standardized sections (SUMMARY, EXPERIENCE, EDUCATION, SKILLS) and consistent bullet formatting.",
        "• **Quantification:** Added/retained measurable impact where context allowed to strengthen credibility."
    ]
    return "\n".join(bullets)


def _generate_explanation(original: str, enhanced: str, job: str) -> str:
    """Generate structured explanation referencing sections & job content."""
    # First attempt AI detailed explanation
    ai_text = None
    if model:
        try:
            prompt = f"""
You are an expert resume coach. Compare ORIGINAL vs ENHANCED for JOB CONTENT.
OUTPUT STRICTLY:
• 5-7 bullets.
• Each bullet: **SECTION / FOCUS:** concise change; include 1 short quoted phrase ("...") added or strengthened from ENHANCED that aligns with the job; finish with why it improves fit.
• Use SECTION names (SUMMARY, EXPERIENCE, EDUCATION, SKILLS, PROJECTS, CERTIFICATIONS) when applicable; if cross-sectional use **OVERALL IMPACT**.
No generic advice, no asking for info.

JOB CONTENT (truncated):\n{job[:3000]}\n---
ORIGINAL (truncated):\n{original[:3500]}\n---
ENHANCED (truncated):\n{enhanced[:3500]}\n---
"""
            resp = model.generate_content(prompt)
            ai_text = (resp.text or '').strip() if resp else ''
        except Exception as e:
            logger.warning(f"AI structured explanation failed: {e}")
            ai_text = None
    if ai_text and len(ai_text.split()) > 12:
        return ai_text
    # Fallback structured deterministic explanation
    return build_deterministic_explanation(original, enhanced, job)

def build_deterministic_explanation(original: str, enhanced: str, job: str) -> str:
    def tokens(s):
        return set(re.findall(r"[A-Za-z]{3,}", s.lower()))
    job_toks = tokens(job)
    orig_sections = parse_sections(original)
    enh_sections = parse_sections(enhanced)
    bullets = []
    for section, eh_lines in enh_sections.items():
        otext = '\n'.join(orig_sections.get(section, []))
        etext = '\n'.join(eh_lines)
        added = (tokens(etext) - tokens(otext)) & job_toks
        if not added:
            continue
        # find example line with an added token
        example = ''
        for l in eh_lines:
            if any(tok in l.lower() for tok in list(added)[:8]):
                example = l.strip()[:140]
                break
        added_list = ', '.join(list(sorted(added))[:5])
        bullets.append(f"• **{section}:** Emphasized {added_list}. e.g. \"{example}\" to mirror role priorities.")
        if len(bullets) >= 6:
            break
    if not bullets:
        return _heuristic_explanation(original, enhanced, job)
    return '\n'.join(bullets)

def parse_sections(text: str) -> dict:
    current = 'GENERAL'
    sections = {current: []}
    for line in text.splitlines():
        l = line.strip()
        if not l:
            continue
        if l.isupper() and 2 <= len(l.split()) <= 5:
            current = l
            sections.setdefault(current, [])
            continue
        sections[current].append(l)
    return sections


def enhance_resume_with_ai(resume_text: str, job_description: str) -> tuple[str, str]:
    """Enhance resume using AI with formatting + explanation fallback."""
    if not model:
        return resume_text, "AI disabled: original resume returned."
    try:
        enhance_prompt = f"""
You are an elite technical resume writer. Rewrite the resume for the job below.
RULES:
- Preserve factual data (companies, titles, dates, degrees).
- Strengthen impact; add conservative metrics ONLY if logically implied.
- Use standard US resume section headers in ALL CAPS: SUMMARY, EXPERIENCE, EDUCATION, SKILLS, (PROJECTS if present), CERTIFICATIONS.
- Bullets start with a strong action verb; no first-person; no pronouns.
- Dense, concise, ATS friendly. First line: candidate name + primary contact placeholders if missing.
- Output ONLY the enhanced resume text. No commentary.

JOB DESCRIPTION (trimmed):\n{job_description[:3000]}\n---
ORIGINAL RESUME (trimmed):\n{resume_text[:5000]}\n---
"""
        raw = model.generate_content(enhance_prompt)
        enhanced_raw = (raw.text or '').strip() if raw else ''
        if not enhanced_raw:
            raise ValueError("Empty AI response")
        normalized = _normalize_ai_text(enhanced_raw)
        sanitized = sanitize_enhanced_content(normalized)
        explanation = _generate_explanation(resume_text, sanitized, job_description)
        return sanitized, explanation
    except Exception as e:
        logger.error(f"Error in AI enhancement: {e}")
        return resume_text, f"AI enhancement temporarily unavailable (fallback). Error: {e}"

#############################
# Job URL scraping utilities
#############################

def is_probable_url(text: str) -> bool:
    return bool(re.match(r'^https?://[^\s]+$', text.strip()))

def fetch_job_content(url: str, timeout: int = 10) -> str:
    """Fetch and extract main textual content from a job posting URL."""
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (JobContentFetcher/1.0)'
        }
        resp = requests.get(url, headers=headers, timeout=timeout)
        if resp.status_code != 200:
            logger.warning(f"Job URL fetch non-200: {resp.status_code}")
            return ''
        soup = BeautifulSoup(resp.text, 'html.parser')
        # Remove script/style/nav/footer
        for tag in soup(['script','style','noscript','header','footer','svg']):
            tag.decompose()
        text_parts = []
        # Domain heuristics
        domain = re.sub(r'^https?://','', url).split('/')[0]
        selectors = []
        if 'greenhouse' in domain:
            selectors += ['div#content', 'div.opening', 'div.main']
        if 'lever.co' in domain:
            selectors += ['div.posting', 'div.content']
        if 'workday' in domain:
            selectors += ['div.body', 'div.GWTCKEditor']
        if 'linkedin.' in domain:
            selectors += ['div.description__text', 'section.description']
        if 'indeed.' in domain:
            selectors += ['div.jobsearch-JobComponent', 'div.jobsearch-JobDescription']
        # Add generic probable containers
        selectors += ['section','article','div']
        grabbed = set()
        for sel in selectors:
            for node in soup.select(sel):
                txt = node.get_text(separator=' ', strip=True)
                if txt and len(txt.split()) > 40 and txt not in grabbed:
                    text_parts.append(txt)
                    grabbed.add(txt)
                    if sum(len(p.split()) for p in text_parts) > 1600:
                        break
            if sum(len(p.split()) for p in text_parts) > 1600:
                break
        if not text_parts:
            # fallback entire body text
            body_txt = soup.get_text(separator=' ', strip=True)
            return ' '.join(body_txt.split()[:1800])
        combined = '\n'.join(text_parts)
        cleaned = re.sub(r'\s+',' ', combined)
        return cleaned[:12000]
    except Exception as e:
        logger.warning(f"Job content fetch failed: {e}")
        return ''

def create_pdf_resume(content: str, output_path: str):
    """Create a professionally formatted PDF with clear hierarchy & wrapping."""
    try:
        from reportlab.lib.enums import TA_CENTER
        # Clean content again for PDF safety
        content = clean_for_pdf(content)

        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=54,
            leftMargin=54,
            topMargin=54,
            bottomMargin=54,
        )

        styles = getSampleStyleSheet()
        base = styles['Normal']
        base.fontSize = 10
        base.leading = 13

        title_style = ParagraphStyle(
            'TITLE', parent=base, fontSize=16, leading=18,
            alignment=TA_CENTER, spaceAfter=10
        )
        section_style = ParagraphStyle(
            'SECTION', parent=base, fontSize=11, leading=14,
            spaceBefore=14, spaceAfter=6, textColor=colors.HexColor('#222222')
        )
        bullet_style = ParagraphStyle(
            'BULLET', parent=base, leftIndent=14, bulletIndent=6,
            spaceBefore=2, spaceAfter=1
        )
        body_style = ParagraphStyle(
            'BODY', parent=base, spaceBefore=1, spaceAfter=4
        )

        lines = [l.rstrip() for l in content.splitlines()]
        first_non_empty = next((l for l in lines if l.strip()), '')
        used_title = False
        story = []

        for raw in lines:
            line = raw.strip()
            if not line:
                story.append(Spacer(1, 4))
                continue
            if not used_title and line == first_non_empty:
                story.append(Paragraph(line, title_style))
                used_title = True
                continue
            if line.isupper() and 2 <= len(line.split()) <= 6 and len(line) <= 48:
                story.append(Paragraph(line, section_style))
                continue
            if line.startswith('- '):
                story.append(Paragraph(line, bullet_style))
                continue
            if len(line) > 180 and '.' in line:
                parts = [p.strip() for p in re.split(r'(?<=\.)\s+', line) if p.strip()]
                for ptxt in parts:
                    story.append(Paragraph(ptxt, body_style))
            else:
                story.append(Paragraph(line, body_style))

        doc.build(story)
    except Exception as e:
        logger.error(f"Error creating PDF: {e}")
        try:
            with open(output_path.replace('.pdf', '.txt'), 'w') as f:
                f.write(content)
        except Exception as fe:
            logger.error(f"Fallback write failed: {fe}")
        raise e

def create_docx_resume(content: str, output_path: str):
    """Generate a DOCX resume mirroring the PDF structure."""
    try:
        doc = Document()
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10.5)
        # Compatibility
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')

        lines = [l.rstrip() for l in clean_for_pdf(content).splitlines()]
        first_non = next((l for l in lines if l.strip()), '')
        used_title = False
        for raw in lines:
            line = raw.strip()
            if not line:
                doc.add_paragraph('')
                continue
            if not used_title and line == first_non:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run.font.size = Pt(16)
                used_title = True
                continue
            if line.isupper() and 2 <= len(line.split()) <= 6 and len(line) <= 48:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
                run.font.size = Pt(11.5)
                continue
            if line.startswith('- '):
                p = doc.add_paragraph(style='List Bullet')
                p.add_run(line[2:])
                continue
            p = doc.add_paragraph(line)
        doc.save(output_path)
    except Exception as e:
        logger.error(f"Error creating DOCX: {e}")
        raise

#############################
# Post-processing sanitation
#############################

def sanitize_enhanced_content(text: str) -> str:
    """Clean AI output: remove markdown asterisks, placeholder brackets, fix headings & bullets."""
    lines = text.splitlines()
    cleaned = []
    header_keywords = {"SUMMARY","EXPERIENCE","EDUCATION","SKILLS","PROJECTS","CERTIFICATIONS","OBJECTIVE","PROFILE"}
    action_verbs = {"developed","designed","implemented","led","managed","optimized","improved","architected","built","created","enhanced","reduced","increased","automated","migrated","refactored"}

    def is_placeholder(l: str) -> bool:
        l_low = l.lower()
        return ('quantifiable achievement' in l_low or 'specify framework' in l_low or 'add another bullet' in l_low or l_low.startswith('[add ') )

    first_content_index = None
    for idx, raw in enumerate(lines):
        line = raw.strip()
        if not line:
            cleaned.append("")
            continue
        # Record first non-empty
        if first_content_index is None:
            first_content_index = idx

        # Pattern like '- *Summary**' or '- *John Smith**' possibly followed by extra asterisks or parenthetical
        m = re.match(r"^- \*([^*]+?)\*+.*$", line)
        if m:
            candidate = m.group(1).strip()
            up = candidate.upper()
            if up in header_keywords:
                cleaned.append(up)
            else:
                # Probably name line; keep original capitalization
                cleaned.append(candidate)
            continue

        # Strip leading '- ' if it appears to be a malformed heading
        if line.startswith('- '):
            after = line[2:].strip()
            # Remove leading/trailing asterisks from after
            after = re.sub(r'^\*+','', after)
            after = re.sub(r'\*+$','', after)
            upper_candidate = after.upper()
            if upper_candidate in header_keywords and len(after.split()) <= 5:
                cleaned.append(upper_candidate)
                continue
            # Placeholder bullet? skip
            if is_placeholder(after):
                continue
            # Real bullet if starts with verb or contains digit/percentage/tech token
            first_word = after.split()[0].lower() if after.split() else ''
            if first_word in action_verbs or re.search(r'(\d|api|sql|flask|python|cloud)', after.lower()):
                cleaned.append(f"- {after}")
                continue
            # Otherwise keep as plain line (maybe part of contact block)
            cleaned.append(after)
            continue

        # Remove surrounding markdown asterisks
        line = re.sub(r'^\*+','', line)
        line = re.sub(r'\*+$','', line)

        # Remove bracket placeholders containing guidance
        line = re.sub(r'\[(?:[^\]]*quantifiable[^\]]*)\]','', line, flags=re.IGNORECASE)
        line = re.sub(r'\[(?:[^\]]*specify[^\]]*)\]','', line, flags=re.IGNORECASE)
        line = re.sub(r'\[(?:[^\]]*add another bullet[^\]]*)\]','', line, flags=re.IGNORECASE)
        line = re.sub(r'\s{2,}',' ', line).strip()
        if not line:
            continue
        # Header detection (non-bullet)
        if (len(line) < 40 and line.upper() in header_keywords):
            cleaned.append(line.upper())
            continue
        cleaned.append(line)

    # Collapse excess blank lines
    final = []
    prev_blank = False
    for l in cleaned:
        if l == "":
            if not prev_blank:
                final.append("")
            prev_blank = True
        else:
            final.append(l)
            prev_blank = False
    return "\n".join(final).strip()

def clean_for_pdf(text: str) -> str:
    """Final cleaning for PDF rendering: remove leftover asterisks, placeholders, malformed bullets."""
    lines = []
    for raw in text.splitlines():
        l = raw.strip()
        if not l:
            lines.append("")
            continue
        # Remove '- *Heading**' patterns
        l = re.sub(r'^- \*([^*]+?)\*+.*$', r'\1', l)
        # Remove stray leading '*'
        l = re.sub(r'^\*(.+)\*$', r'\1', l)
        # Strip placeholder bracket content
        if re.search(r'(quantifiable achievement|add another bullet|specify framework|dates of employment)', l, re.IGNORECASE):
            continue
        # Drop any remaining bracketed guidance completely
        if '[' in l and ']' in l:
            l = re.sub(r'\[[^\]]*\]', '', l).strip()
        # Collapse multiple spaces
        l = re.sub(r'\s{2,}', ' ', l)
        lines.append(l)
    # Ensure first non-empty line not prefixed with '-'
    for i, val in enumerate(lines):
        if val and val.startswith('- '):
            # if likely name/contact (contains @ or digit or |) promote to plain line
            if any(tok in val.lower() for tok in ['@', 'gmail', 'linkedin', ' | ', 'github']):
                lines[i] = val[2:].strip()
            break
    # Remove trailing/leading empty lines
    while lines and not lines[0]:
        lines.pop(0)
    while lines and not lines[-1]:
        lines.pop()
    return "\n".join(lines)

@app.route('/preview/<filename>')
def preview_file(filename):
    """Serve preview files"""
    try:
        file_path = os.path.join(PREVIEWS_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=False)
        else:
            return "File not found", 404
    except Exception as e:
        logger.error(f"Error serving preview: {str(e)}")
        return "Error serving file", 500

@app.route('/download/<filename>')
def download_file(filename):
    """Serve download files"""
    try:
        file_path = os.path.join(PREVIEWS_FOLDER, filename)
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True)
        else:
            return "File not found", 404
    except Exception as e:
        logger.error(f"Error serving download: {str(e)}")
        return "Error serving file", 500

@app.route('/analytics')
def get_analytics():
    """Get basic analytics data"""
    return jsonify({
        'daily_users': analytics_data['daily_users'],
        'total_resumes_generated': analytics_data['total_resumes_generated']
    })

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5001))
    app.run(host='0.0.0.0', port=port, debug=False)
